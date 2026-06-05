"""
Qualizenz Backend Server
Flask application for handling form submissions, analytics, and data management
"""

from flask import Flask, request, jsonify, send_file, send_from_directory, session
import json
import os
import io
import tempfile
from datetime import datetime
from pathlib import Path
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

BASE_DIR = Path(__file__).resolve().parent
app = Flask(__name__, static_folder=None)
app.secret_key = os.environ.get('QUALIZENZ_SECRET_KEY', 'qualizenz-local-owner-secret-change-before-hosting')

# Optional CORS support if flask_cors is installed
try:
    from flask_cors import CORS
    CORS(app)
except ImportError:
    try:
        from flask_cors import CORS
    except (ImportError, ModuleNotFoundError):
        pass
    
    @app.after_request
    def add_cors_headers(response):
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type,Authorization'
        response.headers['Access-Control-Allow-Methods'] = 'GET,POST,PUT,DELETE,OPTIONS'
        return response

# Data storage directory
DATA_DIR = BASE_DIR / 'data'
DATA_DIR.mkdir(exist_ok=True)

# Data file paths
SUBSCRIBERS_FILE = DATA_DIR / 'subscribers.json'
CONTACTS_FILE = DATA_DIR / 'contacts.json'
DOWNLOADS_FILE = DATA_DIR / 'downloads.json'
ANALYTICS_FILE = DATA_DIR / 'analytics.json'
ARTICLES_FILE = DATA_DIR / 'articles.json'
TEMPLATES_FILE = DATA_DIR / 'templates.json'
SITE_SETTINGS_FILE = DATA_DIR / 'site_settings.json'
ADMIN_AUTH_FILE = DATA_DIR / 'admin_auth.json'
DEFAULT_ADMIN_PASSWORD = os.environ.get('QUALIZENZ_ADMIN_PASSWORD', 'Qualizenz@2026')

# Upload storage
UPLOAD_DIR = BASE_DIR / 'uploads'
ARTICLE_UPLOAD_DIR = UPLOAD_DIR / 'articles'
ARTICLE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
ALLOWED_ARTICLE_EXTENSIONS = {'pdf', 'docx'}
MAX_UPLOAD_SIZE = 50 * 1024 * 1024

app.config['MAX_CONTENT_LENGTH'] = MAX_UPLOAD_SIZE

def load_json_file(filepath):
    """Load JSON data from file"""
    if not filepath.exists():
        return []
    try:
        with open(filepath, 'r') as f:
            return json.load(f)
    except json.JSONDecodeError:
        return []

def save_json_file(filepath, data):
    """Save JSON data to file"""
    with open(filepath, 'w') as f:
        json.dump(data, f, indent=2)
    return True

def default_site_settings():
    """Default public display settings managed by the owner dashboard."""
    return {
        'header_banner_text': 'GMP | SOPs | Validation',
        'announcement_text': 'New GMP Audit Readiness Checklist available for download.',
        'announcement_visible': True,
        'hero_title': 'Pharmaceutical GMP articles, SOP templates, validation documents and practical compliance resources.',
        'hero_intro': 'Qualizenz publishes original content and downloadable resources for QA, QC, production, validation, engineering, audit, utilities and pharmaceutical startup teams.',
        'footer_disclaimer': 'Qualizenz provides educational resources, templates, and consulting support. All templates and guidance must be reviewed, customized, approved, and controlled by qualified personnel before use in any GMP-regulated environment.',
        'sections': {
            'headline': True,
            'latestArticles': True,
            'store': True,
            'consulting': True
        },
        'homepage_note': '',
        'updated_at': datetime.now().isoformat()
    }

def load_site_settings():
    """Load site display settings, creating defaults when missing."""
    if not SITE_SETTINGS_FILE.exists():
        settings = default_site_settings()
        save_json_file(SITE_SETTINGS_FILE, settings)
        return settings

    settings = load_json_file(SITE_SETTINGS_FILE)
    if not isinstance(settings, dict):
        settings = default_site_settings()

    defaults = default_site_settings()
    defaults.update(settings)
    defaults['sections'].update(settings.get('sections', {}))
    return defaults

def load_admin_auth():
    """Load admin password hash. Default password is used until the owner changes it."""
    auth = load_json_file(ADMIN_AUTH_FILE)
    if not isinstance(auth, dict):
        auth = {}
    auth.setdefault('password_hash', generate_password_hash(DEFAULT_ADMIN_PASSWORD))
    auth.setdefault('updated_at', datetime.now().isoformat())
    return auth

def save_admin_auth(password):
    auth = {
        'password_hash': generate_password_hash(password),
        'updated_at': datetime.now().isoformat()
    }
    save_json_file(ADMIN_AUTH_FILE, auth)
    return auth

def verify_password(password, stored_hash):
    try:
        return check_password_hash(stored_hash, password)
    except ValueError:
        return False

def is_admin_authenticated():
    return bool(session.get('admin_authenticated'))

def allowed_article_file(filename):
    """Return True when the uploaded article file can be branded for download."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_ARTICLE_EXTENSIONS

def find_article(article_id):
    """Find an article by ID from the JSON article store."""
    articles = load_json_file(ARTICLES_FILE)
    for article in articles:
        if article.get('id') == article_id:
            return article, articles
    return None, articles

def draw_qualizenz_branding(canvas, doc=None):
    """Draw Qualizenz header, footer, watermark and download timestamp."""
    from reportlab.lib.colors import Color, HexColor
    from reportlab.lib.pagesizes import A4

    width, height = A4
    generated_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    canvas.saveState()
    canvas.setFillColor(HexColor('#1D4E89'))
    canvas.setFont('Helvetica-Bold', 12)
    canvas.drawString(42, height - 34, 'Qualizenz')
    canvas.setFillColor(HexColor('#6B7280'))
    canvas.setFont('Helvetica', 8)
    canvas.drawString(42, height - 47, 'Quality | Compliance | Validation | Excellence')

    canvas.setStrokeColor(HexColor('#D1D5DB'))
    canvas.line(42, height - 58, width - 42, height - 58)
    canvas.line(42, 44, width - 42, 44)

    canvas.setFillColor(HexColor('#6B7280'))
    canvas.setFont('Helvetica', 8)
    canvas.drawString(42, 28, f'Downloaded from Qualizenz on {generated_at}')
    canvas.drawRightString(width - 42, 28, 'www.qualizenz.com')

    canvas.translate(width / 2, height / 2)
    canvas.rotate(35)
    canvas.setFillColor(Color(0.12, 0.31, 0.54, alpha=0.08))
    canvas.setFont('Helvetica-Bold', 54)
    canvas.drawCentredString(0, 0, 'QUALIZENZ')
    canvas.restoreState()

def brand_pdf(input_path, output_path):
    """Apply Qualizenz branding to every page of an uploaded PDF."""
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas

    reader = PdfReader(str(input_path))
    writer = PdfWriter()

    for page in reader.pages:
        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)
        packet = io.BytesIO()
        overlay = canvas.Canvas(packet, pagesize=(page_width, page_height))

        generated_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        overlay.saveState()
        overlay.setFillColorRGB(0.11, 0.31, 0.54)
        overlay.setFont('Helvetica-Bold', 12)
        overlay.drawString(36, page_height - 30, 'Qualizenz')
        overlay.setFillColorRGB(0.42, 0.45, 0.50)
        overlay.setFont('Helvetica', 8)
        overlay.drawString(36, page_height - 43, 'Quality | Compliance | Validation | Excellence')
        overlay.line(36, page_height - 54, page_width - 36, page_height - 54)
        overlay.line(36, 42, page_width - 36, 42)
        overlay.drawString(36, 26, f'Downloaded from Qualizenz on {generated_at}')
        overlay.drawRightString(page_width - 36, 26, 'www.qualizenz.com')
        overlay.translate(page_width / 2, page_height / 2)
        overlay.rotate(35)
        overlay.setFillColorRGB(0.11, 0.31, 0.54, alpha=0.08)
        overlay.setFont('Helvetica-Bold', 54)
        overlay.drawCentredString(0, 0, 'QUALIZENZ')
        overlay.restoreState()
        overlay.save()

        packet.seek(0)
        watermark_page = PdfReader(packet).pages[0]
        page.merge_page(watermark_page)
        writer.add_page(page)

    with open(output_path, 'wb') as branded_file:
        writer.write(branded_file)

def convert_docx_to_branded_pdf(input_path, output_path):
    """Create a branded Qualizenz PDF from the text content of an uploaded DOCX file."""
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    document = Document(str(input_path))
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        'QualizenzBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        spaceAfter=8
    )
    title_style = ParagraphStyle(
        'QualizenzTitle',
        parent=styles['Heading1'],
        textColor='#1D4E89',
        fontSize=18,
        leading=22,
        spaceAfter=14
    )

    story = [Paragraph('Qualizenz Article Download', title_style), Spacer(1, 0.15 * inch)]
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            story.append(Paragraph(text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style))

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                story.append(Paragraph(' | '.join(cells).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=42,
        leftMargin=42,
        topMargin=76,
        bottomMargin=70
    )
    doc.build(story, onFirstPage=draw_qualizenz_branding, onLaterPages=draw_qualizenz_branding)

def create_branded_download(input_path, original_name):
    """Return a temporary branded PDF path for PDF or DOCX source files."""
    extension = input_path.suffix.lower()
    output = Path(tempfile.gettempdir()) / f"qualizenz_{input_path.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"

    if extension == '.pdf':
        brand_pdf(input_path, output)
    elif extension == '.docx':
        convert_docx_to_branded_pdf(input_path, output)
    else:
        raise ValueError('Only PDF and DOCX files can be branded for download')

    safe_stem = Path(secure_filename(original_name)).stem or 'qualizenz-article'
    return output, f'{safe_stem}-qualizenz-branded.pdf'

# ==================== NEWSLETTER ENDPOINTS ====================

@app.route('/api/newsletter/subscribe', methods=['POST'])
def subscribe_newsletter():
    """Add email to newsletter subscribers"""
    try:
        data = request.json
        email = data.get('email', '').strip().lower()
        
        if not email or '@' not in email:
            return jsonify({'error': 'Invalid email address'}), 400
        
        subscribers = load_json_file(SUBSCRIBERS_FILE)
        
        if email in subscribers:
            return jsonify({
                'success': False,
                'message': 'Email already subscribed'
            }), 200
        
        subscribers.append(email)
        save_json_file(SUBSCRIBERS_FILE, subscribers)
        
        # Log analytics
        log_event('newsletter_subscribe', {'email': email})
        
        return jsonify({
            'success': True,
            'message': 'Successfully subscribed',
            'total_subscribers': len(subscribers)
        }), 201
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/newsletter/subscribers', methods=['GET'])
def get_subscribers_count():
    """Get total subscriber count"""
    subscribers = load_json_file(SUBSCRIBERS_FILE)
    return jsonify({
        'total_subscribers': len(subscribers),
        'last_updated': datetime.now().isoformat()
    })

# ==================== CONTACT FORM ENDPOINTS ====================

@app.route('/api/contact/submit', methods=['POST'])
def submit_contact():
    """Submit contact form"""
    try:
        data = request.json
        
        contact_record = {
            'name': data.get('name', 'Anonymous').strip(),
            'email': data.get('email', '').strip(),
            'topic': data.get('topic', 'general'),
            'message': data.get('message', '').strip(),
            'submitted_at': datetime.now().isoformat()
        }
        
        if not contact_record['email'] or not contact_record['message']:
            return jsonify({'error': 'Email and message are required'}), 400
        
        contacts = load_json_file(CONTACTS_FILE)
        contacts.append(contact_record)
        save_json_file(CONTACTS_FILE, contacts)
        
        log_event('contact_form_submit', contact_record)
        
        return jsonify({
            'success': True,
            'message': 'Thank you for contacting us',
            'total_contacts': len(contacts)
        }), 201
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contact/list', methods=['GET'])
def get_contacts():
    """Get all contact messages (admin only)"""
    contacts = load_json_file(CONTACTS_FILE)
    return jsonify({
        'total_contacts': len(contacts),
        'contacts': contacts
    })

# ==================== TEMPLATE DOWNLOAD TRACKING ====================

@app.route('/api/templates/download', methods=['POST'])
def track_template_download():
    """Track template downloads"""
    try:
        data = request.json
        template_id = data.get('template_id')
        
        download_record = {
            'template_id': template_id,
            'user_ip': request.remote_addr,
            'downloaded_at': datetime.now().isoformat()
        }
        
        downloads = load_json_file(DOWNLOADS_FILE)
        downloads.append(download_record)
        save_json_file(DOWNLOADS_FILE, downloads)
        
        log_event('template_download', download_record)
        
        return jsonify({
            'success': True,
            'message': 'Download tracked',
            'total_downloads': len(downloads)
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/templates/download-stats', methods=['GET'])
def get_download_stats():
    """Get download statistics by template"""
    downloads = load_json_file(DOWNLOADS_FILE)
    
    stats = {}
    for download in downloads:
        template_id = download.get('template_id')
        stats[template_id] = stats.get(template_id, 0) + 1
    
    return jsonify({
        'total_downloads': len(downloads),
        'templates_stats': stats
    })

# ==================== ANALYTICS ENDPOINTS ====================

def log_event(event_name, event_data=None):
    """Log analytics event"""
    analytics = load_json_file(ANALYTICS_FILE)
    
    event = {
        'event_name': event_name,
        'event_data': event_data or {},
        'timestamp': datetime.now().isoformat(),
        'user_ip': request.remote_addr if request else None
    }
    
    analytics.append(event)
    save_json_file(ANALYTICS_FILE, analytics)

@app.route('/api/analytics/page-view', methods=['POST'])
def track_page_view():
    """Track page views"""
    try:
        data = request.json
        
        log_event('page_view', {
            'page': data.get('page'),
            'referrer': data.get('referrer')
        })
        
        return jsonify({'success': True}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/analytics/dashboard', methods=['GET'])
def get_analytics_dashboard():
    """Get analytics dashboard data"""
    subscribers = load_json_file(SUBSCRIBERS_FILE)
    contacts = load_json_file(CONTACTS_FILE)
    downloads = load_json_file(DOWNLOADS_FILE)
    analytics = load_json_file(ANALYTICS_FILE)
    
    # Count events
    page_views = len([e for e in analytics if e.get('event_name') == 'page_view'])
    
    return jsonify({
        'total_subscribers': len(subscribers),
        'total_contacts': len(contacts),
        'total_downloads': len(downloads),
        'total_page_views': page_views,
        'last_updated': datetime.now().isoformat()
    })

# ==================== ARTICLES ENDPOINTS ====================

@app.route('/api/articles/list', methods=['GET'])
def get_articles():
    """Get all articles"""
    articles = load_json_file(ARTICLES_FILE)
    return jsonify({
        'total_articles': len(articles),
        'articles': articles
    })

@app.route('/api/articles/add', methods=['POST'])
def add_article():
    """Add or update an article draft/published record."""
    try:
        data = request.json
        articles = load_json_file(ARTICLES_FILE)
        article_id = data.get('id') or f"article_{datetime.now().timestamp()}"
        existing_article = next((row for row in articles if row.get('id') == article_id), {})
        
        article = {
            'id': article_id,
            'title': data.get('title'),
            'summary': data.get('summary', ''),
            'content': data.get('content'),
            'category': data.get('category', 'GMP'),
            'tags': data.get('tags', []),
            'author': data.get('author', 'Qualizenz Team'),
            'status': data.get('status', 'Draft'),
            'publish_date': data.get('publishDate') or data.get('publish_date', ''),
            'featured_image': data.get('featuredImage') or data.get('featured_image', ''),
            'seo_title': data.get('seoTitle') or data.get('seo_title', ''),
            'seo_description': data.get('seoDescription') or data.get('seo_description', ''),
            'keywords': data.get('keywords', ''),
            'related_articles': data.get('relatedArticles') or data.get('related_articles', ''),
            'created_at': existing_article.get('created_at', datetime.now().isoformat()),
            'updated_at': datetime.now().isoformat(),
            'files': existing_article.get('files', [])
        }
        
        articles = [row for row in articles if row.get('id') != article_id]
        articles.append(article)
        save_json_file(ARTICLES_FILE, articles)
        
        log_event('article_saved', {'article_id': article['id'], 'status': article['status']})
        
        return jsonify({
            'success': True,
            'message': 'Article saved',
            'article_id': article['id'],
            'total_articles': len(articles)
        }), 201
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/articles/<article_id>/upload-file', methods=['POST'])
def upload_article_file(article_id):
    """Upload a public article PDF or DOCX file for branded download."""
    try:
        article, articles = find_article(article_id)
        if not article:
            return jsonify({'error': 'Article not found'}), 404

        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        uploaded_file = request.files['file']
        if uploaded_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_article_file(uploaded_file.filename):
            return jsonify({'error': 'Only PDF and DOCX files are supported for branded downloads'}), 400

        extension = uploaded_file.filename.rsplit('.', 1)[1].lower()
        original_name = secure_filename(uploaded_file.filename)
        stored_name = secure_filename(f"{article_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.{extension}")
        stored_path = ARTICLE_UPLOAD_DIR / stored_name
        uploaded_file.save(stored_path)

        file_record = {
            'filename': stored_name,
            'original_name': original_name,
            'file_type': extension,
            'uploaded_at': datetime.now().isoformat(),
            'download_count': 0,
            'public_download_url': f'/api/articles/{article_id}/download/{stored_name}'
        }

        article.setdefault('files', []).append(file_record)
        article['updated_at'] = datetime.now().isoformat()
        save_json_file(ARTICLES_FILE, articles)

        return jsonify({
            'success': True,
            'message': 'Article file uploaded. Public download will generate a Qualizenz-branded PDF.',
            'file': file_record
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/articles/<article_id>/download/<filename>', methods=['GET'])
def download_branded_article(article_id, filename):
    """Public download endpoint that returns a Qualizenz-branded PDF."""
    try:
        article, articles = find_article(article_id)
        if not article:
            return jsonify({'error': 'Article not found'}), 404

        safe_filename = secure_filename(filename)
        file_record = None
        for item in article.get('files', []):
            if item.get('filename') == safe_filename:
                file_record = item
                break

        if not file_record:
            return jsonify({'error': 'File not found for this article'}), 404

        source_path = ARTICLE_UPLOAD_DIR / safe_filename
        if not source_path.exists():
            return jsonify({'error': 'Stored file missing'}), 404

        branded_path, download_name = create_branded_download(source_path, file_record.get('original_name', safe_filename))
        file_record['download_count'] = file_record.get('download_count', 0) + 1
        article['updated_at'] = datetime.now().isoformat()
        save_json_file(ARTICLES_FILE, articles)

        downloads = load_json_file(DOWNLOADS_FILE)
        downloads.append({
            'article_id': article_id,
            'filename': safe_filename,
            'branded_download_name': download_name,
            'downloaded_at': datetime.now().isoformat(),
            'user_ip': request.remote_addr
        })
        save_json_file(DOWNLOADS_FILE, downloads)

        return send_file(
            branded_path,
            mimetype='application/pdf',
            as_attachment=False,
            download_name=download_name,
            max_age=0
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==================== TEMPLATES ENDPOINTS ====================

@app.route('/api/templates/list', methods=['GET'])
def get_templates():
    """Get all templates"""
    templates = load_json_file(TEMPLATES_FILE)
    return jsonify({
        'total_templates': len(templates),
        'templates': templates
    })

@app.route('/api/templates/add', methods=['POST'])
def add_template():
    """Add new template"""
    try:
        data = request.json
        
        template = {
            'id': data.get('id', f"template_{datetime.now().timestamp()}"),
            'title': data.get('title'),
            'description': data.get('description'),
            'category': data.get('category', 'QA'),
            'type': data.get('type', 'free'),  # free or premium
            'download_url': data.get('download_url'),
            'created_at': datetime.now().isoformat()
        }
        
        templates = load_json_file(TEMPLATES_FILE)
        templates.append(template)
        save_json_file(TEMPLATES_FILE, templates)
        
        return jsonify({
            'success': True,
            'message': 'Template added',
            'template_id': template['id'],
            'total_templates': len(templates)
        }), 201
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==================== HEALTH & STATUS ENDPOINTS ====================

@app.route('/api/site-settings', methods=['GET'])
def get_site_settings():
    """Get public homepage display settings."""
    return jsonify(load_site_settings())

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    """Authenticate the single owner admin session."""
    data = request.json or {}
    password = data.get('password', '')
    auth = load_admin_auth()
    if not verify_password(password, auth.get('password_hash', '')):
        return jsonify({'success': False, 'error': 'Invalid admin password'}), 401

    session['admin_authenticated'] = True
    return jsonify({'success': True, 'message': 'Admin login successful'}), 200

@app.route('/api/admin/logout', methods=['POST'])
def admin_logout():
    """End the owner admin session."""
    session.pop('admin_authenticated', None)
    return jsonify({'success': True, 'message': 'Admin logged out'}), 200

@app.route('/api/admin/change-password', methods=['POST'])
def admin_change_password():
    """Change the owner admin password."""
    if not is_admin_authenticated():
        return jsonify({'error': 'Admin login required'}), 401

    data = request.json or {}
    current_password = data.get('current_password', '')
    new_password = data.get('new_password', '')
    auth = load_admin_auth()

    if not verify_password(current_password, auth.get('password_hash', '')):
        return jsonify({'error': 'Current password is incorrect'}), 401

    if len(new_password) < 8:
        return jsonify({'error': 'New password must be at least 8 characters'}), 400

    save_admin_auth(new_password)
    return jsonify({'success': True, 'message': 'Admin password updated'}), 200

@app.route('/api/admin/site-settings', methods=['POST'])
def update_site_settings():
    """Update homepage display settings from the owner dashboard."""
    if not is_admin_authenticated():
        return jsonify({'error': 'Admin login required'}), 401

    try:
        data = request.json or {}
        settings = load_site_settings()

        settings['header_banner_text'] = data.get('header_banner_text', settings['header_banner_text']).strip()
        settings['announcement_text'] = data.get('announcement_text', settings.get('announcement_text', '')).strip()
        settings['announcement_visible'] = bool(data.get('announcement_visible', settings.get('announcement_visible', True)))
        settings['hero_title'] = data.get('hero_title', settings['hero_title']).strip()
        settings['hero_intro'] = data.get('hero_intro', settings['hero_intro']).strip()
        settings['homepage_note'] = data.get('homepage_note', settings.get('homepage_note', '')).strip()
        settings['footer_disclaimer'] = data.get('footer_disclaimer', settings.get('footer_disclaimer', '')).strip()

        incoming_sections = data.get('sections', {})
        if isinstance(incoming_sections, dict):
            for key in ['headline', 'latestArticles', 'store', 'consulting']:
                settings['sections'][key] = bool(incoming_sections.get(key, settings['sections'].get(key, True)))

        settings['updated_at'] = datetime.now().isoformat()
        save_json_file(SITE_SETTINGS_FILE, settings)

        return jsonify({
            'success': True,
            'message': 'Site settings updated',
            'settings': settings
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/', methods=['GET'])
def serve_homepage():
    """Serve the Qualizenz website when running through Flask."""
    return send_from_directory(BASE_DIR, 'index.html')

@app.route('/admin/login', methods=['GET'])
@app.route('/admin/dashboard', methods=['GET'])
def serve_admin_dashboard():
    """Serve the owner admin login/dashboard shell."""
    return send_from_directory(BASE_DIR, 'admin_dashboard.html')

@app.route('/<path:filename>', methods=['GET'])
def serve_static_file(filename):
    """Serve static assets and pages when running through Flask."""
    safe_name = secure_filename(filename) if '/' not in filename else filename
    blocked = {'app.py', 'data_manager.py', 'article_manager.py', 'database_schema.sql', 'requirements.txt'}
    if safe_name in blocked or safe_name.startswith('data/') or safe_name.startswith('uploads/'):
        return jsonify({'error': 'File not available'}), 404

    path = BASE_DIR / safe_name
    if path.exists() and path.is_file():
        return send_from_directory(BASE_DIR, safe_name)

    return jsonify({'error': 'File not found'}), 404

@app.route('/api/status', methods=['GET'])
def api_status():
    """API health check"""
    return jsonify({
        'status': 'online',
        'service': 'Qualizenz Backend',
        'version': '1.0.0',
        'timestamp': datetime.now().isoformat()
    })

@app.route('/api/stats', methods=['GET'])
def get_stats():
    """Get comprehensive statistics"""
    subscribers = load_json_file(SUBSCRIBERS_FILE)
    contacts = load_json_file(CONTACTS_FILE)
    downloads = load_json_file(DOWNLOADS_FILE)
    articles = load_json_file(ARTICLES_FILE)
    templates = load_json_file(TEMPLATES_FILE)
    
    return jsonify({
        'subscribers': len(subscribers),
        'contacts': len(contacts),
        'downloads': len(downloads),
        'articles': len(articles),
        'templates': len(templates),
        'timestamp': datetime.now().isoformat()
    })

# ==================== ERROR HANDLERS ====================

@app.errorhandler(404)
def not_found(error):
    """404 error handler"""
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    """500 error handler"""
    return jsonify({'error': 'Internal server error'}), 500

# ==================== MAIN ====================

if __name__ == '__main__':
    print("""
    ╔═══════════════════════════════════════════════════════════╗
    ║      Qualizenz Backend Server                             ║
    ║      Flask API for Content Management & Analytics         ║
    ╚═══════════════════════════════════════════════════════════╝
    
    Server starting on http://localhost:8000
    
    Available endpoints:
    
    Newsletter:
    - POST   /api/newsletter/subscribe
    - GET    /api/newsletter/subscribers
    
    Contacts:
    - POST   /api/contact/submit
    - GET    /api/contact/list
    
    Templates:
    - GET    /api/templates/list
    - POST   /api/templates/add
    - POST   /api/templates/download
    - GET    /api/templates/download-stats
    
    Articles:
    - GET    /api/articles/list
    - POST   /api/articles/add
    
    Analytics:
    - POST   /api/analytics/page-view
    - GET    /api/analytics/dashboard
    
    Status:
    - GET    /api/status
    - GET    /api/stats
    """)
    
    app.run(debug=True, host='0.0.0.0', port=8000)
